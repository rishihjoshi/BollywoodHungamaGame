"""
Update BollywoodHungamaGame.docx to reflect the current app structure and rules.
Run: py scripts/update_docx.py
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCX_PATH = r'C:\Vinayak\Personal\Developement\AI\GIT_Repo\BollywoodHungamaGame\BollywoodHungamaGame.docx'

# ── Build document content ────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    """Add a bold heading paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18 * (indent + 1))
    run = p.add_run(f'• {text}')
    return p

def add_numbered(doc, text, n, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18 * (indent + 1))
    run = p.add_run(f'{n}. {text}')
    return p

def build_doc():
    doc = Document(DOCX_PATH)

    # Clear all existing paragraphs
    for p in doc.paragraphs:
        element = p._element
        element.getparent().remove(element)

    # Clear all tables
    for table in doc.tables:
        element = table._element
        element.getparent().remove(element)

    # ── Title ──
    title = doc.add_paragraph()
    title_run = title.add_run('BollywoodHungamaGame')
    title_run.bold = True
    title_run.font.size = Pt(18)

    doc.add_paragraph()  # spacer

    # ── 1. Overview ──
    add_heading(doc, '1. Overview')
    add_body(doc,
        'BollywoodHungamaGame is a mobile-first Progressive Web App (PWA) that hosts two '
        'Bollywood-themed party games on a shared phone. The app is fully static, works '
        'offline after first load, and requires no backend, authentication, or build step. '
        'It is deployed via GitHub Pages and installable on Android and iOS.'
    )
    add_body(doc,
        'Live URL: https://rishihjoshi.github.io/BollywoodHungamaGame'
    )
    add_body(doc,
        'Tech stack: Vanilla HTML / CSS / JavaScript, Service Worker (PWA), '
        'Playwright E2E test suite. No frameworks, no bundlers.'
    )

    doc.add_paragraph()

    # ── 2. App Structure ──
    add_heading(doc, '2. App Structure')
    add_body(doc, 'The app presents a home hub with two game cards. Each game is self-contained.')

    add_heading(doc, '2.1 Screens — Kaun Hai Gabbar? (KHG)', level=2)
    for i, s in enumerate([
        'Home Hub — game selection',
        'Setup — enter player names (3–12), toggle category hint for Dakus',
        'Role Reveal — hot-seat privacy gate per player; tap to see secret role + word',
        'Offline Play — players give verbal clues and discuss (no UI interaction needed)',
        'End Screen — record whether each Daku was caught; auto-calculates winner',
        'Scores — cumulative session stats (rounds played, Gaon/Daku wins, per-player stats)',
        'Rules — illustrated How to Play overlay',
    ], 1):
        add_numbered(doc, s, i)

    add_heading(doc, '2.2 Screens — Picture Abhi Baaki Hai (PABH)', level=2)
    for i, s in enumerate([
        'Home Hub — game selection',
        'Setup — enter player names (3–12), choose round count (3/5/7), choose timer (60/90/120s)',
        'Combo Screen — pitcher sees their random 4-card combo; cards flip in with stagger animation',
        'Timer Screen — countdown ring + compact combo reference; pitcher pitches their movie idea',
        'Vote Screen — all non-pitchers tap to vote for their favourite pitch',
        'Round Result — winner announced with confetti, standings updated',
        'Final Screen — champion declared, full scoreboard, play again or return home',
    ], 1):
        add_numbered(doc, s, i)

    doc.add_paragraph()

    # ── 3. Objectives ──
    add_heading(doc, '3. Objectives')
    add_bullet(doc, 'Start a game within 60 seconds of opening the link')
    add_bullet(doc, 'Works offline after first load (Service Worker caching)')
    add_bullet(doc, 'Privacy between turns on a shared device (privacy gate screens)')
    add_bullet(doc, 'No backend, no accounts, no runtime dependencies')
    add_bullet(doc, 'Installable as a PWA on Android and iOS')
    add_bullet(doc, 'Supports 3–12 players per game')

    doc.add_paragraph()

    # ── 4. Non-Goals ──
    add_heading(doc, '4. Non-Goals')
    add_bullet(doc, 'No authentication or user accounts')
    add_bullet(doc, 'No real-time multiplayer sync or networking')
    add_bullet(doc, 'No backend, database, or server-side logic')
    add_bullet(doc, 'No payment or monetisation layer')

    doc.add_paragraph()

    # ── 5. How to Play — KHG ──
    add_heading(doc, '5. How to Play — Kaun Hai Gabbar?')

    add_heading(doc, '5.1 Overview', level=2)
    add_body(doc,
        'Kaun Hai Gabbar? is a Sholay-themed social deduction and word-guessing party game. '
        'One or more players are secretly assigned as Daku (Gabbar\'s gang). The rest are Gaon '
        '(villagers) who know the secret Bollywood word. Dakus must bluff their way through '
        'without knowing the word. Gaon players must identify and catch all the Dakus.'
    )

    add_heading(doc, '5.2 Player Count & Daku Distribution', level=2)
    add_bullet(doc, '3–5 players → 1 Daku (Gabbar)')
    add_bullet(doc, '6–8 players → 2 Dakus (Gabbar + Sambha)')
    add_bullet(doc, '9–12 players → 3 Dakus (Gabbar + Sambha + Kaalia)')

    add_heading(doc, '5.3 Characters', level=2)
    add_body(doc, 'Daku team (randomly assigned to Daku players):')
    add_bullet(doc, 'Gabbar 💀 — "Kitne aadmi the?"', indent=1)
    add_bullet(doc, 'Sambha 🔫 — "Poore pachaas hazaar!"', indent=1)
    add_bullet(doc, 'Kaalia ⚔️ — Gabbar\'s loyal foot soldier', indent=1)
    add_body(doc, 'Gaon team (randomly assigned from 14 characters including Veeru, Jay, Thakur, Basanti, and more).')

    add_heading(doc, '5.4 Word Categories', level=2)
    add_bullet(doc, 'Film Title')
    add_bullet(doc, 'Famous Dialogue')
    add_bullet(doc, 'Celebrity')
    add_bullet(doc, 'Song')
    add_bullet(doc, 'Character')

    add_heading(doc, '5.5 Game Flow — Step by Step', level=2)
    add_numbered(doc, 'SETUP — Enter player names (min 3, max 12). '
                 'Optional: toggle whether Dakus see the category hint (on by default). '
                 'Tap "Start Game".', 1)
    add_numbered(doc, 'ROLE REVEAL — The phone is passed to each player in turn. '
                 'Each player sees a privacy gate showing only their name. '
                 'They tap "Reveal My Role" alone (no one else watching). '
                 'Their secret role card flips in showing: team (DAKU/GAON), character name + '
                 'flavour quote, and — for Gaon players — the full secret word, hint, and image. '
                 'Daku players see only the category (e.g. "Film Title") and a reminder to bluff. '
                 'Tap "Done" to lock the card and pass to the next player.', 2)
    add_numbered(doc, 'OFFLINE PLAY — All players have seen their roles. '
                 'Players take turns giving exactly one verbal clue about the secret word. '
                 'Dakus must give a plausible clue without knowing the word. '
                 'Players discuss and debate who the Dakus are. No app interaction needed during this phase.', 3)
    add_numbered(doc, 'END SCREEN — Tap "Done Playing". '
                 'The secret word and all Daku identities are revealed. '
                 'For each Daku, the group decides: was this player caught? Tap YES or NO. '
                 'Results auto-calculate when all Dakus are recorded.', 4)
    add_numbered(doc, 'WIN CONDITION — Gaon wins if ALL Dakus were identified. '
                 'Daku wins if even one Daku escaped detection. '
                 'Confetti fires on Gaon victory.', 5)
    add_numbered(doc, 'PLAY AGAIN — Tap "Play Again" to keep the same players, '
                 're-shuffle roles, and pick a new word. Scores accumulate across rounds.', 6)

    add_heading(doc, '5.6 Scoring', level=2)
    add_body(doc,
        'Lifetime stats stored in localStorage (persist across sessions):'
    )
    add_bullet(doc, 'Total rounds played')
    add_bullet(doc, 'Gaon wins / Daku wins')
    add_bullet(doc, 'Per player: times assigned as Gabbar, times caught as Daku')

    doc.add_paragraph()

    # ── 6. How to Play — PABH ──
    add_heading(doc, '6. How to Play — Picture Abhi Baaki Hai')

    add_heading(doc, '6.1 Overview', level=2)
    add_body(doc,
        'Picture Abhi Baaki Hai is a Bollywood party pitching game. Each round one player '
        'is the Pitcher — they get a random combo of Actor + Location + Genre + Wildcard '
        'and must improvise a Bollywood movie pitch using all four elements. '
        'Everyone else votes for the best pitch. Points go to the most-voted pitcher.'
    )

    add_heading(doc, '6.2 Player Count & Setup Options', level=2)
    add_bullet(doc, '3–12 players')
    add_bullet(doc, 'Round count: 3, 5 (default), or 7 rounds')
    add_bullet(doc, 'Pitch timer: 60s (Fast & chaotic), 90s (Sweet spot — Recommended), 120s (Full filmi drama)')

    add_heading(doc, '6.3 The Combo', level=2)
    add_body(doc, 'Each round the pitcher draws a 4-card combo:')
    add_bullet(doc, 'ACTOR 🎬 — a Bollywood actor name (40-entry pool, no repeats until exhausted)')
    add_bullet(doc, 'LOCATION 📍 — a setting or place (40-entry pool)')
    add_bullet(doc, 'GENRE 🎭 — a film genre (20-entry pool)')
    add_bullet(doc, 'WILDCARD ⚡ — a twist, dialogue, decade, or co-star constraint (40-entry pool)')
    add_body(doc, 'Cards are drawn without replacement; when a pool is exhausted it resets and reshuffles.')

    add_heading(doc, '6.4 Game Flow — Step by Step', level=2)
    add_numbered(doc, 'SETUP — Enter player names, choose rounds and timer. Tap "Start Game".', 1)
    add_numbered(doc, 'COMBO SCREEN — The current pitcher\'s name is shown. '
                 'Four cards flip in one by one with a stagger animation, revealing the combo. '
                 'The Wildcard card has a special glowing border to mark it. '
                 'The pitcher studies the combo. Tap "Start Pitching →" when ready.', 2)
    add_numbered(doc, 'TIMER SCREEN — A circular countdown ring starts. '
                 'The compact combo is shown at the top for reference. '
                 'The ring turns orange below 50% time and red + pulsing below 20%. '
                 'The pitcher improvises their Bollywood movie pitch out loud — '
                 'character, story, drama, songs, all of it. '
                 'When time runs out the app auto-advances to the vote screen. '
                 'Pitcher can also tap "Time\'s Up" to end early.', 3)
    add_numbered(doc, 'VOTE SCREEN — Each non-pitcher player taps their vote for the '
                 'best pitch. Vote category rotates each round (Best Overall Pitch, '
                 'Most Absurd Premise, Would Actually Watch, Funniest Pitch, '
                 'Biggest Blockbuster, Crowd Favourite). '
                 'When all eligible voters have voted, the app auto-advances.', 4)
    add_numbered(doc, 'ROUND RESULT — The round winner is announced with confetti. '
                 'Standings update immediately. Tap "Next Round →" to continue.', 5)
    add_numbered(doc, 'FINAL SCREEN — After all rounds, the champion is declared '
                 'with a random Bollywood subtitle. Full scoreboard is shown. '
                 'Tap "Play Again" to restart with the same players, '
                 'or "Back to Hungama" to return to the home hub.', 6)

    add_heading(doc, '6.5 Scoring', level=2)
    add_bullet(doc, 'Round winner gets points equal to votes received (e.g. 3 votes = +3 points)')
    add_bullet(doc, 'In a tie: all tied players get +1 point each')
    add_bullet(doc, 'If nobody voted: no points awarded that round')
    add_bullet(doc, 'Scores are stored in sessionStorage — they reset when the browser session ends')

    add_heading(doc, '6.6 Pitcher Order', level=2)
    add_body(doc,
        'Players are assigned as pitcher in a shuffled random order at game start. '
        'The order cycles through all players; once everyone has pitched, '
        'the order reshuffles for a new cycle.'
    )

    doc.add_paragraph()

    # ── 7. Technical Design ──
    add_heading(doc, '7. Technical Design')

    add_heading(doc, '7.1 Architecture', level=2)
    add_bullet(doc, 'Static PWA: single index.html, styles.css, app.js, pabh.js, pabh.css')
    add_bullet(doc, 'No framework, no bundler, no build step')
    add_bullet(doc, 'Service Worker (sw.js) caches all assets for offline use')
    add_bullet(doc, 'Hosted on GitHub Pages — zero infrastructure')

    add_heading(doc, '7.2 Data Files', level=2)
    add_bullet(doc, 'data/words.json — KHG word bank (Film Titles, Dialogues, Celebrities, Songs, Characters)')
    add_bullet(doc, 'data/pabh-data.json — PABH pools: 40 actors, 40 locations, 20 genres, 40 wildcards')

    add_heading(doc, '7.3 State Management', level=2)
    add_bullet(doc, 'KHG: localStorage (persists across sessions) for scores; in-memory for game state')
    add_bullet(doc, 'PABH: sessionStorage (resets on browser close) for game state and scores')

    add_heading(doc, '7.4 Privacy System', level=2)
    add_body(doc,
        'KHG uses a mandatory hot-seat privacy gate: each player sees a screen showing only '
        'their name with a "Reveal My Role" button. The role card is only shown after the '
        'player explicitly taps, ensuring no other player sees their secret assignment.'
    )

    add_heading(doc, '7.5 Testing', level=2)
    add_bullet(doc, 'E2E test suite: 28 Playwright tests covering both games end-to-end')
    add_bullet(doc, 'Test runner: npx playwright test')
    add_bullet(doc, 'Static analysis: ESLint (eslint.config.mjs) — 0 errors, 0 warnings')
    add_bullet(doc, 'Dependency audit: npm audit — 0 known CVEs')

    doc.add_paragraph()

    # ── 8. Non-Functional Requirements ──
    add_heading(doc, '8. Non-Functional Requirements')
    add_bullet(doc, 'Works on Chrome, Safari (mobile), Edge')
    add_bullet(doc, 'No console errors in production')
    add_bullet(doc, 'Responsive across common phone sizes (portrait mode)')
    add_bullet(doc, 'Buttons ≥ 44px tap target height')
    add_bullet(doc, 'Accessible font sizes and contrast ratios')
    add_bullet(doc, 'App loads in < 2 seconds on standard mobile network')
    add_bullet(doc, 'Works offline after first load (Service Worker)')

    doc.add_paragraph()

    # ── 9. Future Enhancements ──
    add_heading(doc, '9. Future Enhancements')
    add_bullet(doc, 'KHG: How to Play screen inside PABH (currently only KHG has rules screen)')
    add_bullet(doc, 'PABH: Vote privacy gate — each voter votes secretly without seeing prior votes')
    add_bullet(doc, 'PABH: Abandon game confirmation modal')
    add_bullet(doc, 'Sound effects and haptic feedback')
    add_bullet(doc, 'Shareable round results')
    add_bullet(doc, 'Daily challenge mode (seed-based combos/words)')
    add_bullet(doc, 'TMDB integration for movie posters in KHG word cards')
    add_bullet(doc, 'Third game: Antakshari / SongLink mode')

    doc.add_paragraph()

    # ── 10. Definition of Done ──
    add_heading(doc, '10. Definition of Done')
    add_bullet(doc, 'All E2E acceptance criteria covered by Playwright tests (28 passing)')
    add_bullet(doc, 'Tested with a real 3–5 player group scenario')
    add_bullet(doc, 'No critical UX confusion during gameplay')
    add_bullet(doc, 'Deployed and playable via GitHub Pages')
    add_bullet(doc, 'Service Worker confirmed caching all assets in DevTools → Application → Cache Storage')
    add_bullet(doc, 'ESLint: 0 errors; npm audit: 0 CVEs')

    doc.save(DOCX_PATH)
    print(f'Saved: {DOCX_PATH}')

if __name__ == '__main__':
    build_doc()
